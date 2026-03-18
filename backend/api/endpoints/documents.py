import shutil
import uuid
from importlib import import_module
from io import BytesIO
import mimetypes
from pathlib import Path
from typing import Annotated

import numpy as np
from core.config import settings
from document_loader.format import Format
from document_loader.text_splitter import create_recursive_text_splitter
from entities.document import Document
from fastapi import APIRouter, File, HTTPException, UploadFile
from helpers.log import get_logger
from bs4 import BeautifulSoup
from docx import Document as DocxDocument
from pypdf import PdfReader
from schemas.documents import DocumentInfo, DocumentListResponse, DocumentUploadResponse

from api.deps import VectorDatabaseDep

logger = get_logger(__name__)

router = APIRouter()

# In-memory store of document metadata; keyed by document_id.
_documents: dict[str, DocumentInfo] = {}


def _hydrate_documents_from_disk() -> None:
    """Rebuild in-memory document metadata from docs folder if available."""
    if _documents:
        return

    docs_path = settings.DOCS_PATH
    if not docs_path.exists():
        return

    hydrated: dict[str, DocumentInfo] = {}
    for child in docs_path.iterdir():
        if not child.is_dir():
            continue

        document_id = child.name
        files = [f for f in child.iterdir() if f.is_file()]
        if not files:
            continue

        # Keep deterministic behavior when multiple files exist in one document folder.
        file_path = sorted(files, key=lambda p: p.name)[0]
        guessed_type, _ = mimetypes.guess_type(file_path.name)

        hydrated[document_id] = DocumentInfo(
            document_id=document_id,
            filename=file_path.name,
            size=file_path.stat().st_size,
            content_type=guessed_type or "application/octet-stream",
        )

    _documents.update(hydrated)


def split_chunks(sources: list[Document], chunk_size: int = 1000, chunk_overlap: int = 50) -> list[Document]:
    splitter = create_recursive_text_splitter(
        format=Format.MARKDOWN.value,
        chunk_size=chunk_size,
        chunk_overlap=chunk_overlap,
    )
    return list(splitter.split_documents(sources))


def _extract_pdf_text(content: bytes) -> str:
    """Extract text from PDF using multiple parsers for better compatibility."""
    errors: list[str] = []

    # Basic signature check to fail fast for renamed/non-PDF files.
    if not content.lstrip().startswith(b"%PDF-"):
        raise ValueError("The uploaded file does not appear to be a valid PDF file.")

    # 1) pypdf (already a dependency)
    try:
        reader = PdfReader(BytesIO(content), strict=False)
        pages = [(page.extract_text() or "") for page in reader.pages]
        text = "\n".join(pages).strip()
        if text:
            return text
        errors.append("pypdf extracted no text")
    except Exception as exc:
        errors.append(f"pypdf: {exc}")

    # 2) PyMuPDF fallback if installed
    try:
        fitz = import_module("fitz")
        with fitz.open(stream=content, filetype="pdf") as doc:
            text = "\n".join((page.get_text("text") or "") for page in doc).strip()
        if text:
            return text
        errors.append("pymupdf extracted no text")
    except Exception as exc:
        errors.append(f"pymupdf: {exc}")

    # 3) pdfminer fallback if installed
    try:
        pdfminer_high_level = import_module("pdfminer.high_level")
        text = (pdfminer_high_level.extract_text(BytesIO(content)) or "").strip()
        if text:
            return text
        errors.append("pdfminer extracted no text")
    except Exception as exc:
        errors.append(f"pdfminer: {exc}")

    # 4) OCR fallback for scanned/image-only PDFs
    try:
        fitz = import_module("fitz")
        rapidocr_mod = import_module("rapidocr_onnxruntime")
        ocr_engine = rapidocr_mod.RapidOCR()

        ocr_text_parts: list[str] = []
        with fitz.open(stream=content, filetype="pdf") as doc:
            # Keep OCR bounded for latency/cost on large PDFs.
            max_pages = min(len(doc), max(1, settings.OCR_MAX_PAGES))
            for page_index in range(max_pages):
                page = doc[page_index]
                pix = page.get_pixmap(matrix=fitz.Matrix(settings.OCR_SCALE, settings.OCR_SCALE), alpha=False)
                arr = np.frombuffer(pix.samples, dtype=np.uint8).reshape(pix.height, pix.width, pix.n)

                ocr_result, _ = ocr_engine(arr)
                if not ocr_result:
                    continue

                # rapidocr line format: [box, text, score]
                for line in ocr_result:
                    if len(line) >= 2 and isinstance(line[1], str) and line[1].strip():
                        ocr_text_parts.append(line[1].strip())

                # Stop early once we have enough text for retrieval indexing.
                if sum(len(part) for part in ocr_text_parts) >= settings.OCR_MIN_CHARS:
                    break

        ocr_text = "\n".join(ocr_text_parts).strip()
        if ocr_text:
            return ocr_text
        errors.append("ocr extracted no text")
    except Exception as exc:
        errors.append(f"ocr: {exc}")

    raise ValueError(
        "Unable to read this PDF. It may be corrupted, password-protected, or image-only scanned. "
        f"Parser details: {' | '.join(errors)}"
    )


def extract_text(content: bytes, suffix: str) -> str:
    """Extract plain text from supported file types."""
    if suffix in {".md", ".txt", ".csv", ".json", ".xml"}:
        return content.decode("utf-8")

    if suffix in {".html", ".htm"}:
        html = content.decode("utf-8", errors="ignore")
        soup = BeautifulSoup(html, "html.parser")
        return soup.get_text(separator="\n", strip=True)

    if suffix == ".pdf":
        return _extract_pdf_text(content)

    if suffix == ".docx":
        doc = DocxDocument(BytesIO(content))
        paragraphs = [p.text for p in doc.paragraphs if p.text]
        return "\n".join(paragraphs).strip()

    raise ValueError(f"File type '{suffix}' is not supported for text extraction")


@router.post(
    "/documents",
    response_model=DocumentUploadResponse,
    status_code=201,
    responses={
        400: {"description": "Bad Request - Invalid file type."},
        409: {"description": "Conflict - Document with the same filename already exists."},
    },
)
async def upload_document(file: Annotated[UploadFile, File(...)], index: VectorDatabaseDep):
    """
    Upload a document to the knowledge base.

    Args:
        file: The file to upload. Must have an allowed extension.
        index: Vector database dependency for storing document chunks.

    Returns:
        DocumentUploadResponse containing the generated document_id and filename.

    Raises:
        HTTPException: 400 if file type is not supported.
        HTTPException: 409 if a document with the same filename already exists.
    """
    _hydrate_documents_from_disk()

    suffix = Path(file.filename or "").suffix.lower()
    if suffix not in settings.ALLOWED_UPLOAD_EXTENSIONS:
        raise HTTPException(
            status_code=400,
            detail=f"File type '{suffix}' not supported. Allowed: {sorted(settings.ALLOWED_UPLOAD_EXTENSIONS)}",
        )

    for doc in _documents.values():
        if doc.filename == file.filename:
            raise HTTPException(
                status_code=409,
                detail=f"Document '{file.filename}' already exists.",
            )

    document_id = str(uuid.uuid4())
    dest_dir = settings.DOCS_PATH / document_id
    file_path = dest_dir / (file.filename or document_id)
    content = await file.read()

    try:
        page_content = extract_text(content, suffix)
        if not page_content:
            raise ValueError("No readable text content was found in the uploaded file.")
    except Exception as exc:
        logger.warning(
            "Failed to extract text from uploaded file '%s': %s",
            file.filename,
            exc,
        )
        raise HTTPException(
            status_code=400,
            detail=f"Could not read file content. {exc}",
        )

    dest_dir.mkdir(parents=True, exist_ok=True)
    file_path.write_bytes(content)

    document = Document(
        page_content=page_content,
        metadata={
            "source": str(file_path),
            "document_id": document_id,
            "filename": file.filename,
            "content_type": file.content_type,
            "size": len(content),
        },
    )
    chunks = split_chunks([document], chunk_size=settings.CHUNK_SIZE, chunk_overlap=settings.CHUNK_OVERLAP)
    num_chunks = len(chunks)

    logger.info(f"Number of generated chunks: {num_chunks}")
    logger.info("Adding document chunks to the vector database index...")

    try:
        index.from_chunks(chunks)
    except Exception as exc:
        logger.exception("Failed to index uploaded document '%s': %s", file.filename, exc)
        if dest_dir.exists():
            shutil.rmtree(dest_dir)
        if "OPENAI_API_KEY is not set" in str(exc):
            raise HTTPException(
                status_code=500,
                detail="Backend is missing OPENAI_API_KEY. Set it in .env and restart backend.",
            )
        raise HTTPException(status_code=500, detail="Failed to index uploaded document.")

    logger.info("Memory Index has been updated successfully!")

    doc_info = DocumentInfo(
        document_id=document_id,
        filename=file.filename or document_id,
        size=len(content),
        content_type=file.content_type or "application/octet-stream",
    )
    _documents[document_id] = doc_info

    return DocumentUploadResponse(document_id=document_id, filename=doc_info.filename)


@router.get("/documents", response_model=DocumentListResponse)
async def list_documents():
    """
    List all uploaded documents.

    Returns:
        DocumentListResponse containing a list of all document metadata.
    """
    _hydrate_documents_from_disk()
    return DocumentListResponse(documents=list(_documents.values()))


@router.delete(
    "/documents/{document_id}",
    status_code=204,
    responses={404: {"description": "Not Found - Document with the given ID does not exist."}},
)
async def delete_document(document_id: str, index: VectorDatabaseDep):
    """
    Delete a document from the knowledge base.

    Removes the document's metadata, associated file from disk, and should remove
    chunks from the vector database index (currently not fully implemented).

    Args:
        document_id: The unique identifier of the document to delete.
        index: Vector database dependency for removing document chunks.

    Raises:
        HTTPException: 404 if the document with the given ID is not found.
    """
    _hydrate_documents_from_disk()

    if document_id not in _documents:
        raise HTTPException(status_code=404, detail=f"Document '{document_id}' not found.")

    dest_dir = settings.DOCS_PATH / document_id
    if dest_dir.exists():
        shutil.rmtree(dest_dir)

    # TODO: implement an efficient way to remove all chunks associated with this document from the vector database index
    # https://github.com/umbertogriffo/rag-chatbot/pull/10#discussion_r2936567674

    del _documents[document_id]
